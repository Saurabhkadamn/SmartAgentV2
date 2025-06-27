# utils/loaders.py
import io
import PyPDF2
import docx
import pandas as pd
from typing import List, Union
import streamlit as st

def load_file(uploaded_file) -> List[str]:
    """
    Load and extract text from various file types.
    
    Args:
        uploaded_file: Streamlit uploaded file object
        
    Returns:
        List of strings containing the extracted text
    """
    file_type = uploaded_file.type
    file_content = uploaded_file.read()
    
    try:
        if file_type == "application/pdf":
            return load_pdf(file_content)
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return load_docx(file_content)
        elif file_type in ["application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", 
                          "application/vnd.ms-excel"]:
            return load_xlsx(file_content)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
            
    except Exception as e:
        st.error(f"Error loading file: {str(e)}")
        return []

def load_pdf(file_content: bytes) -> List[str]:
    """Extract text from PDF file."""
    text_pages = []
    pdf_file = io.BytesIO(file_content)
    
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                text = page.extract_text()
                if text.strip():  # Only add non-empty pages
                    text_pages.append(f"Page {page_num + 1}:\n{text}")
            except Exception as e:
                st.warning(f"Could not extract text from page {page_num + 1}: {str(e)}")
                
    except Exception as e:
        raise Exception(f"Error reading PDF: {str(e)}")
    
    return text_pages

def load_docx(file_content: bytes) -> List[str]:
    """Extract text from DOCX file."""
    try:
        docx_file = io.BytesIO(file_content)
        doc = docx.Document(docx_file)
        
        paragraphs = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                paragraphs.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                if row_text.strip():
                    paragraphs.append(row_text)
        
        return ["\n".join(paragraphs)] if paragraphs else []
        
    except Exception as e:
        raise Exception(f"Error reading DOCX: {str(e)}")

def load_xlsx(file_content: bytes) -> List[str]:
    """Extract text from Excel file."""
    try:
        excel_file = io.BytesIO(file_content)
        
        # Read all sheets
        excel_data = pd.read_excel(excel_file, sheet_name=None, engine='openpyxl')
        
        text_content = []
        
        for sheet_name, df in excel_data.items():
            # Convert DataFrame to string representation
            sheet_text = f"Sheet: {sheet_name}\n\n"
            
            # Add column headers
            headers = " | ".join(df.columns.astype(str))
            sheet_text += f"Headers: {headers}\n\n"
            
            # Add data rows
            for index, row in df.iterrows():
                row_text = " | ".join(row.astype(str).fillna(""))
                sheet_text += f"Row {index + 1}: {row_text}\n"
            
            text_content.append(sheet_text)
        
        return text_content
        
    except Exception as e:
        raise Exception(f"Error reading Excel file: {str(e)}")

# Additional utility functions
def clean_text(text: str) -> str:
    """Clean and normalize extracted text."""
    import re
    
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    
    # Remove special characters but keep basic punctuation
    text = re.sub(r'[^\w\s\.\,\!\?\;\:\-\(\)]', '', text)
    
    return text.strip()

def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
    """Split text into overlapping chunks."""
    chunks = []
    start = 0
    
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        
        # Try to break at sentence boundaries
        if end < len(text):
            last_period = chunk.rfind('.')
            last_newline = chunk.rfind('\n')
            break_point = max(last_period, last_newline)
            
            if break_point > start + chunk_size // 2:
                chunk = text[start:break_point + 1]
                end = break_point + 1
        
        chunks.append(chunk.strip())
        start = end - overlap
        
        if start >= len(text):
            break
            
    return [chunk for chunk in chunks if chunk]  # Remove empty chunks